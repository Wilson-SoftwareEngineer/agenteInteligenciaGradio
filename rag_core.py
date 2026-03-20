import os
import io
import uuid
import re
from dotenv import load_dotenv
from qdrant_client import QdrantClient, models
from langchain_qdrant import QdrantVectorStore
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_text_splitters import RecursiveCharacterTextSplitter
import pypdf

try:
    from docx import Document as DocxDocument
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    import pytesseract
    from PIL import Image
    _OCR_OK = True
except ImportError:
    _OCR_OK = False

load_dotenv()

QDRANT_HOST     = os.getenv("QDRANT_HOST", "localhost")
QDRANT_PORT     = int(os.getenv("QDRANT_PORT", 6334))
QDRANT_URL      = f"http://{QDRANT_HOST}:{QDRANT_PORT}"
COLLECTION_NAME = os.getenv("COLLECTION_NAME", "relatorios_ficco")
OPENAI_API_KEY  = os.getenv("OPENAI_API_KEY")
OPENAI_MODEL    = os.getenv("OPENAI_MODEL", "gpt-4o")
EMBEDDINGS_PATH = os.getenv(
    "EMBEDDINGS_PATH",
    os.path.join(os.path.dirname(__file__), "modelo_embeddings"),
)

_qdrant_client   = None
_embeddings      = None
_vector_store    = None
_llm             = None


# ── Conexões (mesmo padrão do agente original) ────────────────────────────────

def get_embeddings() -> HuggingFaceEmbeddings:
    global _embeddings
    if _embeddings is None:
        _embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
            cache_folder=EMBEDDINGS_PATH,
            model_kwargs={"device": "cpu"},
            encode_kwargs={"normalize_embeddings": False},
        )
    return _embeddings


def get_qdrant() -> QdrantClient:
    global _qdrant_client
    if _qdrant_client is None:
        _qdrant_client = QdrantClient(url=QDRANT_URL)
    return _qdrant_client


def get_vector_store() -> QdrantVectorStore:
    global _vector_store
    if _vector_store is None:
        _vector_store = QdrantVectorStore(
            client=get_qdrant(),
            collection_name=COLLECTION_NAME,
            embedding=get_embeddings(),
            content_payload_key="text",
            metadata_payload_key="metadata",
        )
    return _vector_store


def get_llm() -> ChatOpenAI:
    global _llm
    if _llm is None:
        _llm = ChatOpenAI(
            model=OPENAI_MODEL,
            api_key=OPENAI_API_KEY,
            temperature=0,
            streaming=True,
        )
    return _llm


# ── Status do sistema ─────────────────────────────────────────────────────────

def get_status() -> dict:
    try:
        client = get_qdrant()
        info   = client.get_collection(COLLECTION_NAME)
        return {
            "status":       "online",
            "documentos":   info.points_count,
            "dimensoes":    info.config.params.vectors.size,
            "colecao":      COLLECTION_NAME,
            "modelo_llm":   OPENAI_MODEL,
            "porta_qdrant": QDRANT_PORT,
        }
    except Exception as e:
        return {"status": "erro", "mensagem": str(e)}


# ── Busca híbrida (scroll por keyword + similaridade vetorial) ────────────────
# Mesmo padrão do agente original

def search(query: str, top_k: int = 20) -> list:
    client       = get_qdrant()
    vector_store = get_vector_store()

    # Limpar query e extrair termos
    query_limpa = re.sub(r'[?.,;:/]', ' ', query)
    termos = [t.strip() for t in query_limpa.split() if len(t) >= 3]

    contexto = []
    vistos   = set()
    fontes   = set()

    # 1. Busca por keyword via scroll
    if termos:
        filtros = [
            models.FieldCondition(key="text", match=models.MatchText(text=t))
            for t in termos
        ]
        try:
            res_scroll, _ = client.scroll(
                collection_name=COLLECTION_NAME,
                scroll_filter=models.Filter(should=filtros),
                limit=20,
                with_payload=True,
            )
            for res in res_scroll:
                txt = res.payload.get("text", "")
                meta = res.payload.get("metadata", {})
                arq  = meta.get("arquivo", "desconhecido")
                pag  = meta.get("pagina", "-")
                if txt and txt not in vistos:
                    contexto.append({"texto": txt, "fonte": arq, "pagina": pag, "score": 1.0})
                    vistos.add(txt)
                    fontes.add(arq)
        except Exception:
            pass

    # 2. Busca vetorial por similaridade com score real
    try:
        res_vetor = vector_store.similarity_search_with_score(query, k=top_k)
        for doc, score in res_vetor:
            txt = doc.page_content
            arq = doc.metadata.get("arquivo", "desconhecido")
            pag = doc.metadata.get("pagina", "-")
            if txt and txt not in vistos:
                contexto.append({"texto": txt, "fonte": arq, "pagina": pag, "score": round(float(score), 4)})
                vistos.add(txt)
                fontes.add(arq)
    except Exception:
        pass

    return contexto


def busca_avancada(cpf: str = "", nome: str = "", vulgo: str = "", endereco: str = "") -> list:
    client   = get_qdrant()
    contexto = []
    vistos   = set()

    # Filtro exato por CPF — prioridade máxima
    if cpf:
        cpf_limpo = re.sub(r'[.\-]', '', cpf).strip()
        filtros_cpf = [
            models.FieldCondition(key="text", match=models.MatchText(text=cpf)),
            models.FieldCondition(key="text", match=models.MatchText(text=cpf_limpo)),
        ]
        try:
            res_cpf, _ = client.scroll(
                collection_name=COLLECTION_NAME,
                scroll_filter=models.Filter(should=filtros_cpf),
                limit=20,
                with_payload=True,
            )
            for res in res_cpf:
                txt  = res.payload.get("text", "")
                meta = res.payload.get("metadata", {})
                arq  = meta.get("arquivo", "desconhecido")
                pag  = meta.get("pagina", "-")
                if txt and txt not in vistos:
                    contexto.append({"texto": txt, "fonte": arq, "pagina": pag, "score": 1.0})
                    vistos.add(txt)
        except Exception:
            pass

    # Busca híbrida pelos demais campos
    partes = []
    if cpf:      partes.append(f"CPF {cpf}")
    if nome:     partes.append(nome)
    if vulgo:    partes.append(f"vulgo {vulgo}")
    if endereco: partes.append(endereco)
    if not partes:
        return []

    for item in search(" ".join(partes), top_k=20):
        if item["texto"] not in vistos:
            contexto.append(item)
            vistos.add(item["texto"])

    return contexto


# ── Prompt investigativo (adaptado do original) ───────────────────────────────

PROMPT = ChatPromptTemplate.from_messages([
    ("system", """Você é um Analista de Inteligência Sênior da FICCO-PI.
Sua missão é extrair e consolidar todas as informações disponíveis sobre o alvo a partir dos trechos fornecidos.

REGRAS:
1. CONSOLIDAÇÃO TOTAL: Leia TODOS os trechos fornecidos e reúna cada dado encontrado, mesmo que esteja em trechos diferentes.
2. DADOS PESSOAIS: Procure em todos os trechos por CPF (formato 000.000.000-00 ou somente números), data de nascimento, nome da mãe, endereço, RG, telefone e qualquer outro dado pessoal. Se encontrar em qualquer trecho, inclua na resposta.
3. FOCO NO ALVO: Inclua informações de terceiros apenas se forem comparsas diretos, parentes ou vínculos operacionais.
4. INTEGRIDADE: Se um campo não aparecer em nenhum trecho, escreva "Não localizado nos relatórios." — nunca invente dados.
5. SEM ALUCINAÇÃO: Use apenas o que estiver textualmente presente nos trechos.
6. FONTE OBRIGATÓRIA: Toda informação deve citar o relatório de origem.

FORMATO DE RESPOSTA:
📌 RESUMO DE INTELIGÊNCIA
─────────────────────────
👤 ALVO: [NOME COMPLETO]
🆔 CPF: [número encontrado ou "Não localizado nos relatórios."]
🪪 RG: [número encontrado ou "Não localizado nos relatórios."]
🎂 NASCIMENTO: [data encontrada ou "Não localizado nos relatórios."]
👩 MÃE: [nome encontrado ou "Não localizado nos relatórios."]
🏠 ENDEREÇO: [endereço encontrado ou "Não localizado nos relatórios."]
📞 TELEFONE: [número encontrado ou "Não localizado nos relatórios."]
🚗 VÍNCULOS: [placas, comparsas, associados — com fonte]

📂 HISTÓRICO OPERACIONAL:
• [Fato 1 — Fonte: nome do relatório, pág. X]
• [Fato 2 — Fonte: nome do relatório, pág. X]
• [Continue para todos os fatos relevantes encontrados]

⚠️ NOTAS OPERACIONAIS: [análise consolidada com base em todos os trechos]
─────────────────────────
Fonte: Base Integrada FICCO-PI

Trechos recuperados dos relatórios:
{context}"""),
    ("human", "{question}"),
])


def _formatar_contexto(docs: list) -> str:
    return "\n---\n".join(d["texto"] for d in docs)


# ── Chat com streaming ────────────────────────────────────────────────────────

def chat_stream(pergunta: str, historico: list):
    docs     = search(pergunta, top_k=20)
    contexto = _formatar_contexto(docs)
    chain    = PROMPT | get_llm() | StrOutputParser()

    acumulado = ""
    for chunk in chain.stream({"context": contexto, "question": pergunta}):
        acumulado += chunk
        yield acumulado


# ── Consulta avançada com síntese LLM ────────────────────────────────────────

def consulta_avancada_stream(cpf: str = "", nome: str = "", vulgo: str = "", endereco: str = ""):
    docs = busca_avancada(cpf=cpf, nome=nome, vulgo=vulgo, endereco=endereco)
    if not docs:
        yield "Nenhum documento encontrado para os critérios informados."
        return

    contexto = _formatar_contexto(docs)

    # Monta pergunta focada nos campos preenchidos
    partes = []
    if cpf:      partes.append(f"CPF {cpf}")
    if nome:     partes.append(f"nome {nome}")
    if vulgo:    partes.append(f"vulgo {vulgo}")
    if endereco: partes.append(f"endereço {endereco}")
    pergunta = "Forneça o relatório completo de inteligência sobre o alvo com " + ", ".join(partes) + ". Extraia todos os dados pessoais e histórico operacional disponíveis nos trechos."

    chain = PROMPT | get_llm() | StrOutputParser()

    acumulado = ""
    for chunk in chain.stream({"context": contexto, "question": pergunta}):
        acumulado += chunk
        yield acumulado


# ── Ingestão de documentos (PDF e DOCX) ──────────────────────────────────────

def _garantir_colecao():
    """Cria a coleção no Qdrant se ainda não existir."""
    client   = get_qdrant()
    colecoes = [c.name for c in client.get_collections().collections]
    if COLLECTION_NAME not in colecoes:
        client.create_collection(
            collection_name=COLLECTION_NAME,
            vectors_config=models.VectorParams(size=384, distance=models.Distance.COSINE),
        )


def _extrair_paginas_pdf(caminho: str) -> list:
    """Retorna lista de (numero_pagina, texto) para cada página do PDF."""
    reader  = pypdf.PdfReader(caminho)
    paginas = []
    for i, page in enumerate(reader.pages):
        texto = (page.extract_text() or "").strip()
        if texto:
            paginas.append((i + 1, texto))
    return paginas


def _extrair_texto_docx(caminho: str) -> str:
    """Extrai texto de corpo, cabeçalho, rodapé e faz OCR em imagens do DOCX."""
    if not _DOCX_OK:
        raise ImportError("python-docx não está instalado. Execute: pip install python-docx")
    doc  = DocxDocument(caminho)
    texto = ""
    for section in doc.sections:
        for para in section.header.paragraphs:
            if para.text.strip():
                texto += f"[CABEÇALHO]: {para.text}\n"
        for para in section.footer.paragraphs:
            if para.text.strip():
                texto += f"[RODAPÉ]: {para.text}\n"
    for para in doc.paragraphs:
        if para.text.strip():
            texto += para.text + "\n"
    if _OCR_OK:
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    imagem    = Image.open(io.BytesIO(rel.target_part.blob))
                    texto_ocr = pytesseract.image_to_string(imagem, lang="por")
                    if texto_ocr.strip():
                        texto += f"\n[IMAGEM]:\n{texto_ocr}\n"
                except Exception:
                    continue
    return texto


def ingerir_documento(caminho_arquivo: str, progress_callback=None) -> dict:
    """Processa PDF ou DOCX, vetoriza e indexa no Qdrant com IDs determinísticos."""
    try:
        _garantir_colecao()

        embeddings   = get_embeddings()
        client       = get_qdrant()
        nome_arquivo = os.path.basename(caminho_arquivo)
        ext          = os.path.splitext(nome_arquivo)[1].lower()
        splitter     = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=200)

        # ── Extração de texto por formato ────────────────────────────────────
        if ext == ".pdf":
            paginas = _extrair_paginas_pdf(caminho_arquivo)
        elif ext == ".docx":
            texto_docx = _extrair_texto_docx(caminho_arquivo)
            paginas    = [(1, texto_docx)] if texto_docx.strip() else []
        else:
            return {"sucesso": False, "mensagem": f"Formato não suportado: {ext}. Use PDF ou DOCX."}

        if not paginas:
            return {"sucesso": False, "mensagem": "Nenhum texto extraído do documento."}

        # ── Chunking, vetorização e montagem dos pontos ───────────────────────
        pontos = []
        bloco  = 0
        total  = len(paginas)

        for idx, (pagina, texto) in enumerate(paginas):
            if progress_callback:
                progress_callback(idx / total, desc=f"Página {idx + 1}/{total}")

            for chunk in splitter.split_text(texto):
                if len(chunk.strip()) < 50:
                    continue

                vetor     = embeddings.embed_query(chunk)
                string_id = f"{nome_arquivo}_{pagina}_{bloco}"
                point_id  = str(uuid.uuid5(uuid.NAMESPACE_DNS, string_id))
                bloco    += 1

                pontos.append(models.PointStruct(
                    id=point_id,
                    vector=vetor,
                    payload={
                        "text":     chunk,
                        "metadata": {"arquivo": nome_arquivo, "pagina": pagina},
                    },
                ))

        if not pontos:
            return {"sucesso": False, "mensagem": "Nenhum chunk válido gerado."}

        # ── Envio ao Qdrant em batches de 50 ─────────────────────────────────
        for k in range(0, len(pontos), 50):
            client.upsert(collection_name=COLLECTION_NAME, points=pontos[k: k + 50])

        return {
            "sucesso":  True,
            "arquivo":  nome_arquivo,
            "paginas":  total,
            "chunks":   len(pontos),
            "mensagem": f"'{nome_arquivo}' importado. {len(pontos)} chunks indexados.",
        }

    except Exception as e:
        return {"sucesso": False, "mensagem": f"Erro: {str(e)}"}


# ── Teste direto ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Testando conexão com Qdrant...")
    print(get_status())